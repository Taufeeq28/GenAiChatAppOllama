import streamlit as st
from langchain_groq import ChatGroq
import docx
from pathlib import Path
import sys
from typing import List, Dict
import json

class ResumeOptimizer:
    def __init__(self, groq_api_key: str, model_id: str):
        self.llm = ChatGroq(
            groq_api_key=groq_api_key,
            model_name=model_id
        )
        self.added_skills = []
        self.added_projects = []

    def analyze_job_description(self, job_description: str) -> Dict:
        """Extract key requirements from job description"""
        prompt = f"""
        Analyze this job description and extract:
        1. Required technical skills
        2. Experience level
        3. Key responsibilities

        Job Description:
        {job_description}

        Return the analysis in JSON format with keys: 'technical_skills', 'experience_level', 'responsibilities'
        """
        
        response = self.llm.invoke(prompt)
        try:
            return json.loads(response.content)
        except:
            return {
                "technical_skills": [],
                "experience_level": "Not specified",
                "responsibilities": []
            }

    def generate_optimized_projects(self, skills: List[str], experience_level: str) -> List[Dict]:
        """Generate relevant projects based on required skills"""
        prompt = f"""
        Generate 2-3 professional projects that demonstrate expertise in these skills: {', '.join(skills)}
        For someone at {experience_level} level.
        
        Return in JSON format with array of projects, each having:
        - title
        - description
        - technologies (list)
        - responsibilities (list)
        - outcomes (list)
        """
        
        response = self.llm.invoke(prompt)
        try:
            return json.loads(response.content)["projects"]
        except:
            return []

    def process_resume(self, input_path: str, job_description: str, output_path: str) -> bool:
        """Process the resume and create optimized version"""
        try:
            # Load the resume
            doc = docx.Document(input_path)
            
            # Analyze job requirements
            requirements = self.analyze_job_description(job_description)
            self.added_skills = requirements['technical_skills']
            
            # Generate new projects
            new_projects = self.generate_optimized_projects(
                requirements['technical_skills'],
                requirements['experience_level']
            )
            self.added_projects = new_projects
            
            # Find the projects section
            projects_start = -1
            for i, para in enumerate(doc.paragraphs):
                if "PROJECTS" in para.text.upper() or "PROJECT EXPERIENCE" in para.text.upper():
                    projects_start = i
                    break
            
            # If no projects section found, create one
            if projects_start == -1:
                projects_start = len(doc.paragraphs)
                doc.add_paragraph("PROJECTS").style = 'Heading 1'
            
            # Add new projects
            for project in new_projects:
                p = doc.add_paragraph()
                p.add_run(f"{project['title']}\n").bold = True
                p.add_run(project['description'] + "\n")
                p.add_run("Technologies: ").bold = True
                p.add_run(", ".join(project['technologies']) + "\n")
                p.add_run("Key Responsibilities:\n").italic = True
                for resp in project['responsibilities']:
                    p.add_run(f"• {resp}\n")
                p.add_run("Outcomes:\n").italic = True
                for outcome in project['outcomes']:
                    p.add_run(f"• {outcome}\n")
                doc.add_paragraph()  # Add spacing
            
            # Save the optimized resume
            doc.save(output_path)
            return True
            
        except Exception as e:
            st.error(f"Error processing resume: {str(e)}")
            return False

    def extract_sections(self, doc) -> Dict[str, List[str]]:
        """Extract sections from the resume"""
        sections = {
            "header": [],
            "summary": [],
            "experience": [],
            "projects": [],
            "skills": [],
            "education": [],
            "other": []
        }
        
        current_section = "header"
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Identify sections
            if any(keyword in text.lower() for keyword in ["experience", "work history"]):
                current_section = "experience"
            elif "project" in text.lower():
                current_section = "projects"
            elif any(keyword in text.lower() for keyword in ["skills", "technologies"]):
                current_section = "skills"
            elif "education" in text.lower():
                current_section = "education"
            elif any(keyword in text.lower() for keyword in ["summary", "objective"]):
                current_section = "summary"
            else:
                sections[current_section].append(text)
                
        return sections